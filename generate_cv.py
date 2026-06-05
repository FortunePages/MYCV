from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Australian Format CV
def create_australian_cv():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("ROLYN KIUNISALA")
    run.font.size = Pt(24)
    run.font.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Full Stack Software Developer | Software Engineer")
    run.font.size = Pt(11)
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("Cebu, Philippines | rolynkiunisala@gmail.com | +63 906 454 4620\n")
    contact.add_run("LinkedIn: linkedin.com/in/rolyn-kiunisala-b862ab377")
    contact.paragraph_format.space_after = Pt(12)
    
    # Professional Summary
    heading = doc.add_heading('PROFESSIONAL SUMMARY', level=2)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)
    summary_text = "Full Stack Software Developer with over 3 years of experience in designing, developing, and maintaining enterprise and web-based applications. Skilled in backend and frontend development, API integration, and system enhancement within Agile environments. Experienced in delivering client-based solutions, automation initiatives, and cloud-based deployments. Strong ability to adapt quickly to new technologies and contribute to scalable software systems."
    doc.add_paragraph(summary_text)
    
    # Core Skills
    heading = doc.add_heading('CORE SKILLS', level=2)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)
    
    skills = [
        "Programming Languages: Java, JavaScript, C#, Python, PHP",
        "Frontend: HTML, CSS/SCSS, JavaScript, Angular, React (basic)",
        "Backend: Node.js, .NET, Java, Python",
        "Database: MySQL, SQL",
        "Cloud & DevOps: AWS, Azure DevOps, CI/CD, Lambda",
        "Tools: Git, GitHub, VS Code, Visual Studio, GitHub Copilot",
        "Methodologies: Agile / Scrum",
        "Other: REST API Development, System Integration, Debugging, Deployment, Network Basics, SCM"
    ]
    
    for skill in skills:
        doc.add_paragraph(skill, style='List Bullet')
    
    # Professional Experience
    heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)
    
    # Job 1
    job1 = doc.add_paragraph()
    run = job1.add_run("Full Stack Software Developer")
    run.font.bold = True
    job1_meta = doc.add_paragraph("ALMA Phil Manpower Services Corp. – Cebu City, Philippines | March 18, 2026 – Present (Project-Based Contract)")
    job1_meta.paragraph_format.space_after = Pt(3)
    
    job1_details = [
        "Design, develop, and maintain full stack web applications based on client requirements.",
        "Develop and integrate REST APIs and backend services using Node.js, Java, and .NET.",
        "Build responsive frontend interfaces using Angular and modern web technologies.",
        "Perform debugging, troubleshooting, and performance optimization.",
        "Collaborate with project managers and technical teams in Agile environment.",
        "Participate in system enhancements, testing, and deployment activities.",
        "Ensure code quality, scalability, and security compliance."
    ]
    for detail in job1_details:
        doc.add_paragraph(detail, style='List Bullet')
    
    # Job 2
    job2 = doc.add_paragraph()
    run = job2.add_run("Software Engineering / Packaged App Development Associate")
    run.font.bold = True
    job2_meta = doc.add_paragraph("Accenture Inc. – Ebloc 2, Cebu IT Park, Cebu City, Philippines | February 27, 2023 – March 18, 2026")
    job2_meta.paragraph_format.space_after = Pt(3)
    
    job2_details = [
        "Developed and supported enterprise-level software applications.",
        "Worked on packaged application development and system enhancements.",
        "Participated in software testing, debugging, and production support.",
        "Collaborated with global teams in Agile delivery environment.",
        "Delivered system fixes and enhancements within SLA timelines.",
        "Supported API integrations and backend services using Java, .NET, and Node.js."
    ]
    for detail in job2_details:
        doc.add_paragraph(detail, style='List Bullet')
    
    # Education
    heading = doc.add_heading('EDUCATION', level=2)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)
    
    edu1 = doc.add_paragraph()
    run = edu1.add_run("Bachelor of Science in Computer Engineering")
    run.font.bold = True
    doc.add_paragraph("University of Cebu – Main Campus, Sanciangko St., Cebu City | 2017 – 2018")
    
    edu2 = doc.add_paragraph()
    run = edu2.add_run("Arcelo Memorial National High School")
    run.font.bold = True
    doc.add_paragraph("San Vicente, Liloan, Cebu")
    
    edu3 = doc.add_paragraph()
    run = edu3.add_run("Simeon Ayuda Elementary School")
    run.font.bold = True
    doc.add_paragraph("San Vicente, Liloan, Cebu")
    
    # Projects & Achievements
    heading = doc.add_heading('PROJECTS & ACHIEVEMENTS', level=2)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)
    
    achievement = doc.add_paragraph()
    run = achievement.add_run("GEN AI Mini-Hackathon Winner")
    run.font.bold = True
    doc.add_paragraph("Team project, 2025")
    
    achievement_details = [
        "Built a conversational AI agent for structured change request intake.",
        "Collected resources, schedule, environment, and requirement details.",
        "Generated submission-ready structured summaries."
    ]
    for detail in achievement_details:
        doc.add_paragraph(detail, style='List Bullet')
    
    doc.save('Rolyn_Kiunisala_CV.docx')
    print("✅ Australian format CV saved: Rolyn_Kiunisala_CV.docx")

# Professional Format CV
def create_professional_cv():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("ROLYN KIUNISALA")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Full Stack Software Developer")
    run.font.size = Pt(11)
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("Cebu, Philippines | rolynkiunisala@gmail.com | +63 906 454 4620\n")
    contact.add_run("LinkedIn: linkedin.com/in/rolyn-kiunisala-b862ab377")
    contact.paragraph_format.space_after = Pt(12)
    
    # Professional Summary
    heading = doc.add_heading('SUMMARY', level=2)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)
    summary_text = "Full Stack Software Developer with over 3 years of experience in designing, developing, and maintaining enterprise and web-based applications. Skilled in backend and frontend development, API integration, and system enhancement within Agile environments. Experienced in delivering client-based solutions, automation initiatives, and cloud-based deployments. Strong ability to adapt quickly to new technologies and contribute to scalable software systems."
    doc.add_paragraph(summary_text)
    
    # Skills
    heading = doc.add_heading('SKILLS', level=2)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)
    
    skills = [
        "Programming Languages: Java, JavaScript, C#, Python, PHP",
        "Frontend: HTML, CSS/SCSS, JavaScript, Angular, React (basic)",
        "Backend: Node.js, .NET, Java, Python",
        "Database: MySQL, SQL",
        "Cloud & DevOps: AWS, Azure DevOps, CI/CD, Lambda",
        "Tools: Git, GitHub, VS Code, Visual Studio, GitHub Copilot",
        "Methodologies: Agile / Scrum",
        "Other: REST API Development, System Integration, Debugging, Deployment, Network Basics, SCM"
    ]
    
    for skill in skills:
        doc.add_paragraph(skill, style='List Bullet')
    
    # Professional Experience
    heading = doc.add_heading('EXPERIENCE', level=2)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)
    
    # Job 1
    job1 = doc.add_paragraph()
    run = job1.add_run("Full Stack Software Developer")
    run.font.bold = True
    job1_meta = doc.add_paragraph("ALMA Phil Manpower Services Corp. – Cebu City, Philippines | March 18, 2026 – Present (Project-Based Contract)")
    job1_meta.paragraph_format.space_after = Pt(3)
    
    job1_details = [
        "Design, develop, and maintain full stack web applications based on client requirements.",
        "Develop and integrate REST APIs and backend services using Node.js, Java, and .NET.",
        "Build responsive frontend interfaces using Angular and modern web technologies.",
        "Perform debugging, troubleshooting, and performance optimization.",
        "Collaborate with project managers and technical teams in Agile environment.",
        "Participate in system enhancements, testing, and deployment activities.",
        "Ensure code quality, scalability, and security compliance."
    ]
    for detail in job1_details:
        doc.add_paragraph(detail, style='List Bullet')
    
    # Job 2
    job2 = doc.add_paragraph()
    run = job2.add_run("Software Engineering / Packaged App Development Associate")
    run.font.bold = True
    job2_meta = doc.add_paragraph("Accenture Inc. – Ebloc 2, Cebu IT Park, Cebu City, Philippines | February 27, 2023 – March 18, 2026")
    job2_meta.paragraph_format.space_after = Pt(3)
    
    job2_details = [
        "Developed and supported enterprise-level software applications.",
        "Worked on packaged application development and system enhancements.",
        "Participated in software testing, debugging, and production support.",
        "Collaborated with global teams in Agile delivery environment.",
        "Delivered system fixes and enhancements within SLA timelines.",
        "Supported API integrations and backend services using Java, .NET, and Node.js."
    ]
    for detail in job2_details:
        doc.add_paragraph(detail, style='List Bullet')
    
    # Education
    heading = doc.add_heading('EDUCATION', level=2)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)
    
    edu1 = doc.add_paragraph()
    run = edu1.add_run("Bachelor of Science in Computer Engineering")
    run.font.bold = True
    doc.add_paragraph("University of Cebu – Main Campus, Sanciangko St., Cebu City | 2017 – 2018")
    
    edu2 = doc.add_paragraph()
    run = edu2.add_run("Arcelo Memorial National High School")
    run.font.bold = True
    doc.add_paragraph("San Vicente, Liloan, Cebu")
    
    edu3 = doc.add_paragraph()
    run = edu3.add_run("Simeon Ayuda Elementary School")
    run.font.bold = True
    doc.add_paragraph("San Vicente, Liloan, Cebu")
    
    # Projects & Achievements
    heading = doc.add_heading('PROJECTS & ACHIEVEMENTS', level=2)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)
    
    achievement = doc.add_paragraph()
    run = achievement.add_run("GEN AI Mini-Hackathon Winner")
    run.font.bold = True
    doc.add_paragraph("Team project, 2025")
    
    achievement_details = [
        "Built a conversational AI agent for structured change request intake.",
        "Collected resources, schedule, environment, and requirement details.",
        "Generated submission-ready structured summaries."
    ]
    for detail in achievement_details:
        doc.add_paragraph(detail, style='List Bullet')
    
    doc.save('Rolyn_Kiunisala_Professional_CV.docx')
    print("✅ Professional format CV saved: Rolyn_Kiunisala_Professional_CV.docx")

if __name__ == "__main__":
    create_australian_cv()
    create_professional_cv()
    print("\n✅ Both CV documents have been successfully updated!")
